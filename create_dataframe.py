import os
import re
from decimal import Decimal
from functools import lru_cache
from pathlib import Path

import pandas as pd
from docx import Document
import osmnx as ox
from dataclasses import dataclass

from pdf2docx import Converter
import wikipedia

pd.set_option('display.max_colwidth', 200)


@dataclass
class TableStruct:
    project_name: int = None
    grant_name: int = None
    year: int = None
    organization: int = None
    organization_short_name: str = None
    lat: float = None
    lon: float = None


def convert_files() -> list:
    added_files = []
    path_docx = Path('processed_project')

    for pdf_file in Path('projects').glob('*.pdf'):
        filename_docx = Path(pdf_file.stem + '.docx')

        docx_file = path_docx / filename_docx

        if not os.path.exists(docx_file):
            added_files.append(docx_file)

            cv = Converter(str(pdf_file))
            cv.convert(str(docx_file))
            cv.close()

    return added_files


def get_geocode_osm(organization: str) -> tuple[float, float]:
    geocode = ox.geocode_to_gdf(organization)
    return float(geocode.lat.iloc[0]), float(geocode.lon.iloc[0])


def get_geocode_wiki(organization: str) -> tuple[Decimal, Decimal] | tuple[None, None]:
    if not organization:
        return None, None

    wikipedia.set_lang("ru")
    pages = wikipedia.search(organization, results=5)

    for page in pages:
        try:
            return wikipedia.page(page).coordinates
        except (KeyError, wikipedia.exceptions.DisambiguationError, wikipedia.exceptions.PageError):
            continue

    return None, None


@lru_cache
def get_geocode(organization: str):
    try:
        return get_geocode_osm(organization)
    except ValueError:
        return get_geocode_wiki(organization)


def delete_quotes(organization: str) -> str:
    quotes = ('«', '»', '‹', '›', '‟', '”', '\'', '„', '‘', '’')
    for quot in quotes:
        organization = organization.replace(quot, '"')

    if not organization.count('"'):
        return organization

    if organization.count('"') % 2 == 0 and organization[0] == organization[-1] == '"':
        return organization[1:-1]

    if organization.count('"') == 3 and organization[0] == '"':
        return organization[1:]

    return organization


@lru_cache
def get_short_name(organization: str) -> str:
    if organization.upper() == organization:
        organization = organization.capitalize()
    search_short_name = re.search(r'(?<=образования).+|(?<=науки).+|(?<=предприятие).+|(?<=фонд).+',
                                  organization)

    if not search_short_name or len(search_short_name.group(0)) < 5:
        search_short_name = re.search(
            r'(?<=[уУ]чреждение).+|(?<=организация).+|(?<=ответственностью).+|(?<=общество).+',
            organization
        )

        if search_short_name and 'Концерн' in search_short_name.group(0):
            search_short_name = re.search('(?<=Концерн).+', search_short_name[0])

    if search_short_name and len(search_short_name[0]) > 5:
        organization = delete_quotes(search_short_name[0].strip())

        if organization[0].lower() == organization[0]:
            return organization.capitalize()

        return organization

    return organization.strip()


def process_document(document: Document, main_df: pd.DataFrame) -> pd.DataFrame:
    table_struct = TableStruct()

    try:
        year = re.search(r'20\d{2}', document.paragraphs[1].text)
    except IndexError:
        return main_df

    if year:
        table_struct.year = year[0]
    table_struct.grant_name = re.search(r'(?<=«)[^»]*', document.paragraphs[1].text)[0]

    for index, cell in enumerate(document.tables[0].rows[0].cells):
        if re.search(r'Название', cell.text):
            table_struct.project_name = index
        elif re.search(r'Российская организация|Организация', cell.text):
            table_struct.organization = index

    df = pd.DataFrame(columns=[*table_struct.__dict__.keys()])

    for table in document.tables:
        for row in table.rows:
            organization = row.cells[table_struct.organization].text.replace('\n', '')
            project_name = row.cells[table_struct.project_name].text.replace('\n', '')

            if not row.cells[0].text:
                last_row = df.iloc[[df.shape[0] - 1]]

                project_name = last_row.project_name.to_string(index=False) + project_name
                organization = last_row.organization.to_string(index=False) + organization

                df.drop(labels=[df.shape[0] - 1], axis=0)

            short_name = get_short_name(organization)
            lat, lon = get_geocode(short_name)

            current_row = pd.DataFrame(
                dict(
                    zip(
                        table_struct.__dict__.keys(),
                        [
                            project_name,
                            table_struct.grant_name,
                            table_struct.year,
                            organization,
                            short_name,
                            lat,
                            lon
                        ]
                    )
                ),
                index=[0]
            )
            df = pd.concat([df, current_row], ignore_index=True)

    df.drop(labels=[0], axis=0, inplace=True)
    return pd.concat([df, main_df], ignore_index=True)


def create_dataframe():
    if not os.path.exists('grants.csv'):
        main_df = pd.DataFrame(columns=[*TableStruct().__dict__.keys()])
    else:
        main_df = pd.read_csv('grants.csv')

    for document in convert_files():
        main_df = process_document(Document(document), main_df)

    main_df.to_csv(r'grants.csv', index=False)

